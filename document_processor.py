import os
import re
import pypdf
from docx import Document as DocxDocument
from typing import List, Dict, Any
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles processing of PDF, DOCX, and TXT documents"""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)

                if len(pdf_reader.pages) == 0:
                    raise ValueError("PDF file has no pages")

                text = ""
                pages_with_text = 0
                total_pages = len(pdf_reader.pages)

                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                            pages_with_text += 1
                        else:
                            logger.warning(f"Page {i + 1} returned empty text")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {i + 1}: {str(e)}")
                        continue

                text = text.strip()
                if not text:
                    raise ValueError(f"No text could be extracted from PDF. Extracted from {pages_with_text}/{total_pages} pages. The file might be scanned images, have security restrictions, or be corrupted.")

                logger.info(f"Extracted {len(text)} characters from {pages_with_text}/{total_pages} PDF pages")
                return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            raise

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"

            text = text.strip()
            if not text:
                raise ValueError("No text could be extracted from DOCX file")

            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read().strip()

            if not text:
                raise ValueError("TXT file is empty")

            logger.info(f"Extracted {len(text)} characters from TXT file")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {str(e)}")
            raise

    def process_document(self, file_path: str) -> List[LangChainDocument]:
        """Process document and return list of LangChain documents"""
        try:
            ext = os.path.splitext(file_path)[1].lower()

            if ext == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif ext == '.docx':
                text = self.extract_text_from_docx(file_path)
            elif ext == '.txt':
                text = self.extract_text_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {ext}")

            if not text.strip():
                raise ValueError("No text content extracted from document")

            # Normalize whitespace
            text = re.sub(r'\n{2,}', '\n\n', text)

            if len(text) < 10:
                raise ValueError(f"Extracted text is too short ({len(text)} characters). Document may be empty or corrupted.")

            metadata = {
                "source": os.path.basename(file_path),
                "file_path": file_path,
                "file_type": ext,
                "total_length": len(text),
                "word_count": len(text.split())
            }

            chunks = self.text_splitter.split_text(text)
            if not chunks:
                raise ValueError("No chunks created from extracted text")

            documents = []
            for i, chunk in enumerate(chunks):
                if chunk.strip():
                    doc = LangChainDocument(
                        page_content=chunk,
                        metadata={**metadata, "chunk_id": i, "chunk_size": len(chunk)}
                    )
                    documents.append(doc)

            if not documents:
                raise ValueError("All chunks were empty after filtering")

            logger.info(f"{len(documents)} chunks created from {len(text)} characters")
            return documents

        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise

    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """Return summary info about the document before chunking"""
        try:
            ext = os.path.splitext(file_path)[1].lower()

            if ext == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif ext == '.docx':
                text = self.extract_text_from_docx(file_path)
            elif ext == '.txt':
                text = self.extract_text_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {ext}")

            normalized_text = re.sub(r'\n{2,}', '\n\n', text)
            chunks = self.text_splitter.split_text(normalized_text)

            return {
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "file_type": ext,
                "total_length": len(text),
                "word_count": len(text.split()),
                "actual_chunks": len(chunks)
            }

        except Exception as e:
            logger.error(f"Error getting document info for {file_path}: {str(e)}")
            raise
